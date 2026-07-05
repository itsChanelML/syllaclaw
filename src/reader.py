"""
SyllaClaw file reader.
Extracts text from PDF, DOCX, and TXT files.
Returns plain text for NIM to parse.
"""

import os
from pathlib import Path
from typing import Optional


SUPPORTED = {".pdf", ".docx", ".doc", ".txt", ".md"}


def extract_text(file_path: Path) -> Optional[str]:
    """
    Extract text from a syllabus file.
    Returns None if file is unreadable (triggers ESCALATE).
    """
    suffix = file_path.suffix.lower()

    if suffix not in SUPPORTED:
        return None

    if suffix == ".txt" or suffix == ".md":
        return _read_text(file_path)

    if suffix == ".pdf":
        return _read_pdf(file_path)

    if suffix in (".docx", ".doc"):
        return _read_docx(file_path)

    return None


def _read_text(path: Path) -> Optional[str]:
    try:
        content = path.read_text(encoding="utf-8", errors="ignore").strip()
        return content if len(content) > 50 else None
    except Exception:
        return None


def _read_pdf(path: Path) -> Optional[str]:
    try:
        from pdfminer.high_level import extract_text as pdfminer_extract
        text = pdfminer_extract(str(path))
        if not text or len(text.strip()) < 50:
            return None
        return text.strip()
    except ImportError:
        # Fallback — try reading as raw text (won't work for scanned PDFs)
        try:
            content = path.read_bytes()
            # Very naive PDF text extraction
            import re
            text = content.decode("latin-1", errors="ignore")
            # Extract text between BT and ET markers
            matches = re.findall(r'BT.*?ET', text, re.DOTALL)
            if matches:
                raw = " ".join(matches)
                cleaned = re.sub(r'[^\x20-\x7E\n]', ' ', raw)
                cleaned = re.sub(r'\s+', ' ', cleaned).strip()
                return cleaned if len(cleaned) > 50 else None
        except Exception:
            return None
    except Exception:
        return None


def _read_docx(path: Path) -> Optional[str]:
    try:
        from docx import Document
        doc = Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also get table text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        text = "\n".join(paragraphs).strip()
        return text if len(text) > 50 else None
    except ImportError:
        # python-docx not installed — try reading as zip/xml
        try:
            import zipfile
            import re
            with zipfile.ZipFile(str(path)) as z:
                with z.open("word/document.xml") as f:
                    xml = f.read().decode("utf-8")
            text = re.sub(r'<[^>]+>', ' ', xml)
            text = re.sub(r'\s+', ' ', text).strip()
            return text if len(text) > 50 else None
        except Exception:
            return None
    except Exception:
        return None


def get_syllabus_files(folder: Path) -> list:
    """
    Return all supported syllabus files in a folder.
    Excludes hidden files, sample files named BROKEN_, etc.
    """
    if not folder.exists():
        return []

    files = []
    for f in sorted(folder.iterdir()):
        if f.name.startswith("."):
            continue
        if f.suffix.lower() in SUPPORTED:
            files.append(f)

    return files